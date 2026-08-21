from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('output/docs/remalj-carewell-school-portal-documentation.docx')
BLUE = '1265B0'
DARK = '17365D'
LIGHT = 'EAF3FB'
INK = '1F2937'
MUTED = '5B6675'

def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None:
        mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}'); mar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')

def set_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')

def borders(table, color='D5DFEB'):
    tblPr = table._tbl.tblPr; b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'5'); el.set(qn('w:color'), color); b.append(el)
    tblPr.append(b)

def add_para(doc, text='', size=11, bold=False, color=INK, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.10
    if align is not None: p.alignment = align
    r = p.add_run(text); set_font(r, size, bold, color, italic)
    return p

def heading(doc, number, title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6); p.paragraph_format.keep_with_next = True
    r = p.add_run(f'{number}.  {title.upper()}'); set_font(r, 13, True, DARK)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(text), 10.5)
    return p

def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False; borders(t)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader'); repeat.set(qn('w:val'), 'true'); trPr.append(repeat)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; set_width(cell, widths[i]); shade(cell, LIGHT); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_font(p.add_run(h), 9.5, True, DARK)
    for row in rows:
        cells=t.add_row().cells
        for i, val in enumerate(row):
            set_width(cells[i], widths[i]); set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05; set_font(p.add_run(val), 9.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def footer(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run('REMALJ CAREWELL INSPIRATIONAL SCHOOL  |  School Portal Documentation'); set_font(r, 8.5, True, BLUE)
    p.add_run('                                      ')
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); p._p.append(fld)

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.85); sec.left_margin = sec.right_margin = Inches(0.85); sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)
footer(sec)
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); styles['Normal'].font.size=Pt(11)

# Cover
add_para(doc, 'REMALJ CAREWELL INSPIRATIONAL SCHOOL', 12, True, BLUE, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
logo = Path('public/remalj-carewell-logo.jpg')
if logo.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(logo), width=Inches(1.25)); p.paragraph_format.space_after=Pt(14)
add_para(doc, 'SCHOOL PORTAL\nDOCUMENTATION', 23, True, BLUE, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'Functional Reference and Implementation Roadmap', 13, False, MUTED, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
rule=doc.add_paragraph(); rule.paragraph_format.space_after=Pt(18)
pPr=rule._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:color'),BLUE); pb.append(bottom); pPr.append(pb)
add_para(doc, 'Documentation purpose', 11, True, DARK, after=4)
add_para(doc, 'This document records the current REMALJ Carewell school portal prototype and translates the attached RCIS Terms of Reference into a practical implementation roadmap. It distinguishes what is demonstrable in the present front-end from capability that still requires production services, policy design, and acceptance testing.', 11, False, INK, after=16)
add_para(doc, 'Project motto: One School. One System. One Source of Truth.', 12, True, DARK, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
meta = add_table(doc, ['DOCUMENT', 'STATUS'], [('School Portal Documentation', 'Prototype reference - implementation roadmap included'), ('Prepared for', 'REMALJ Carewell Inspirational School'), ('Version', '1.0 | 20 August 2026')], [3000, 6360])
doc.add_page_break()

heading(doc, 1, 'Background')
add_para(doc, 'REMALJ Carewell Inspirational School requires a coherent digital environment for teaching, learning, parent engagement, student information, communication and school transport. The existing portal provides a visual, role-based starting point across Teacher, Parent and Student experiences. It is a Vite and React single-page prototype with locally managed sample data and no live back-end services.', 11)
heading(doc, 2, 'Purpose of the Documentation')
add_para(doc, 'The documentation is intended to guide stakeholders, developers and future implementation partners. It describes the visible portal behaviour, identifies boundaries of the current prototype, and sets minimum production requirements aligned to the structure and institutional-control principles of the reference ToR.', 11)
heading(doc, 3, 'Current Solution Architecture')
add_para(doc, 'The current solution is a front-end application. Users switch among Teacher, Parent and Student portals through a shared top bar. Teacher and Parent experiences display a sign-in screen before portal access; Student access is configured as open for demonstration. State is held in the browser session and shared portal data is provided through a React context.', 11)
add_table(doc, ['LAYER', 'CURRENT IMPLEMENTATION', 'PRODUCTION DIRECTION'], [
 ('Presentation', 'React 18 portal views, responsive CSS, Lucide icons and branded assets.', 'Maintain responsive role-based web interface and accessibility controls.'),
 ('Application state', 'React useState and PortalStore context with representative data.', 'Replace with authenticated API services and server-side business rules.'),
 ('Authentication', 'Demo email/password and school-card interface; browser camera scan support where available.', 'Identity provider, role-based access, secure sessions, audit logging and recovery controls.'),
 ('Data and reporting', 'In-browser sample records; report-download helper can create a CSV-style file.', 'Institution-owned database, controlled imports/exports, backups and governed reporting.')
], [1550, 3800, 4010])
heading(doc, 4, 'Portal Access and Roles')
add_table(doc, ['ROLE', 'CURRENT ACCESS', 'PRIMARY FUNCTIONS'], [
 ('Teacher', 'Login screen then teacher dashboard.', 'Class summary, student ledger, activity feed, subject performance, messages and live transport view.'),
 ('Parent / Guardian', 'Login screen then parent dashboard.', 'Child selector, academic results, attendance, fees summary, events, teacher updates, contacts, reports and bus tracking.'),
 ('Student', 'Open demonstration portal.', 'Dashboard, weekly schedule, assignments, deadlines, achievements, quick resources and read-only bus status.'),
 ('Administrator', 'Not yet implemented as a portal role.', 'Required for production configuration, enrolment, fees, user administration, safeguarding, reports and operational control.')
], [1500, 2800, 5060])

heading(doc, 5, 'Functional Scope')
add_para(doc, 'The following capabilities are visible in the prototype or represented as navigation. A page labelled “Coming Soon” is a planned area and must not be treated as implemented functionality.', 11, False, INK, after=5)
add_table(doc, ['MODULE', 'CURRENT STATE', 'NOTES'], [
 ('Authentication and school card', 'Demonstration interface', 'Email/password and card-ID routes; optional camera scan uses browser BarcodeDetector when supported. No credential verification or secure persistence.'),
 ('Academic management', 'Partially represented', 'Teacher and parent dashboards show example grades, performance and schedules. No grade-entry workflow or academic database.'),
 ('Assignments and learning', 'Partially represented', 'Student assignment status and submission action are visual only; no file submission or teacher marking workflow.'),
 ('Parent communication', 'Partially represented', 'Teacher updates, messages, directory and public pages are presented in the interface; delivery and consent controls require a back end.'),
 ('Transport and fleet', 'Interactive prototype', 'Teacher, parent and student views surface bus route/status information. GPS integration, route management and notifications are required for production.'),
 ('Reports', 'Prototype download', 'Published-report helper downloads sample CSV data. Formal report generation and permissions remain to be implemented.'),
 ('Fees, contacts, settings', 'Navigation / representative cards', 'Portal surfaces and some components exist; payment, directory governance and configuration services are not integrated.')
], [1900, 2150, 4310])

heading(doc, 6, 'Required Production Modules')
add_para(doc, 'Following the reference ToR, a production Digital Campus Management Ecosystem should be delivered in phased, testable modules. The present portal can become the user-facing layer for the modules below.', 11)
for item in [
 'Student information and admissions: applicant, learner, guardian, class, attendance, transfers, graduation, documents and alumni records.',
 'Academic management: curriculum, subjects, timetables, lesson planning, assessment, examinations, report cards, progress tracking and approved parent access.',
 'Finance and administration: fee structures, invoicing, receipting, balances, scholarships, payment integration, cash controls and financial reporting.',
 'Human resources and payroll: staff records, recruitment, contracts, attendance, leave, payroll, statutory deductions and appraisal records.',
 'Boarding, dining, library, STEM, studio, clubs, facilities and asset management: operational records connected to a single learner/staff identity.',
 'Health, welfare and safeguarding: restricted-access case management, first-aid records, emergency contacts, escalation and anonymised management analytics.',
 'Management command centre: authorised dashboards for enrolment, attendance, academic outcomes, fees, transport, safeguarding, assets and service performance.'
]: bullet(doc, item)

heading(doc, 7, 'Data Ownership, Security and Safeguarding')
add_para(doc, 'All institutional data must remain under REMALJ Carewell control. Any implementation partner must provide exportable data, documented interfaces and administration credentials. Production design must enforce least-privilege role-based access, multi-factor authentication for privileged accounts, encryption in transit and at rest, audit trails, controlled backups, patch management, monitoring and tested incident-response arrangements.', 11)
add_para(doc, 'Health, welfare, safeguarding and pickup information require heightened controls. Access must be limited to authorised personnel, information sharing must follow institutional policy and applicable law, and all actions must be auditable.', 11, False, INK, after=8)

heading(doc, 8, 'Implementation Approach')
add_table(doc, ['PHASE', 'OUTCOME', 'ACCEPTANCE EVIDENCE'], [
 ('I - Foundation', 'Institution-owned hosting, central data model, identity, admissions, core academic records and portal authentication.', 'Approved architecture, data model, access matrix, core workflows and migrated/validated pilot data.'),
 ('II - Campus operations', 'Transport, fees, communications, boarding/dining, HR, facilities and operational dashboards.', 'Role tests, integration tests, training records and operational reports.'),
 ('III - Future-ready learning', 'Digital learning, library, STEM, studio, portfolios and co-curricular management.', 'Configured modules, user guides and stakeholder acceptance.'),
 ('IV - Intelligence and optimisation', 'Executive analytics, automated alerts, resilience improvements and performance monitoring.', 'Dashboard validation, backup/recovery exercise and support handover.')
], [1600, 4200, 3560])
heading(doc, 9, 'Testing, Documentation and Training')
add_para(doc, 'Installation alone shall not constitute acceptance. Before commissioning, the implementation must complete functional, role/access, integration, data-migration, reporting, performance, security and backup/recovery testing. User acceptance testing shall include teachers, parents/guardians, students, school management, finance, boarding/dining, transport, safeguarding and administrators as appropriate.', 11)
add_para(doc, 'The final delivery set shall include system architecture, installation and configuration records, administrator and user manuals, API/integration references, database documentation, backup/recovery guide, cybersecurity configuration, troubleshooting guide, change log, training materials and a commissioning report.', 11)
heading(doc, 10, 'Known Prototype Limitations')
for item in [
 'No live server, database, identity provider, payment gateway, GPS provider or external API is connected.',
 'Displayed people, academic scores, bus locations, dates, fees and messages are representative demonstration data.',
 'Teacher and Parent sign-in is a front-end state change, not real authentication; Student access is open in the current application configuration.',
 'Several navigation destinations intentionally show a “Coming Soon” view and do not provide complete workflows.',
 'The camera/card scan experience depends on browser support and does not currently verify a scanned card against an institutional registry.'
]: bullet(doc, item)
heading(doc, 11, 'Success Criteria')
add_para(doc, 'The programme will be successful when REMALJ Carewell operates a secure, integrated, documented, maintainable and scalable digital campus environment that it can control without permanent dependency on an implementing vendor. The deployed system must provide trusted records, appropriate access controls, practical training, verified recovery capability and measurable support for teaching, learning, administration, safeguarding and campus operations.', 12, True, DARK, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
